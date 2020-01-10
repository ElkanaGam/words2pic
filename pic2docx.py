from PIL import Image
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH 

img_list = [ './chair.jpg']
for i in img_list:
    im = Image.open(i)
    im.save(i, 'JPEG')


document = Document()
tables = document.tables
table = document.add_table(rows=3, cols=3)
#row_cells = table.add_row().cells


for i, image in enumerate([ './chair.jpg','./chair.jpg', './chair.jpg', './chair.jpg', './chair.jpg', './chair.jpg', './chair.jpg', './chair.jpg', './chair.jpg']):
    print(i)
    row_cells = table.rows[i//3]
    paragraph = row_cells.cells[i%3].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(image , width=Inches(1.7))
    run  = paragraph.add_run()
    run.text = "chair"
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
document.save('doc.docx')